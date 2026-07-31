from pathlib import Path
from typing import Literal

from docx import Document as DocxDocument
from langchain_core.tools import tool
from pypdf import PdfReader

from .config import WORKSPACE_ROOT

MAX_CHARS = 12000
TEXT_EXTENSIONS = {".txt", ".md", ".py", ".csv", ".json", ".rst"}


def _resolve(path: str) -> Path:
    candidate = Path(path)
    candidate = candidate.resolve() if candidate.is_absolute() else (WORKSPACE_ROOT / candidate).resolve()
    if not candidate.is_relative_to(WORKSPACE_ROOT):
        raise ValueError(f"'{path}' is outside the allowed workspace {WORKSPACE_ROOT}")
    return candidate


@tool
def search_files(query: str, mode: Literal["name", "content"] = "name", max_results: int = 30) -> str:
    """Search for files inside the sandboxed workspace.

    mode="name" matches the query as a case-insensitive substring against file names.
    mode="content" greps text-based file contents for the query string.
    Returns paths relative to the workspace root, one per line.
    """
    results: list[str] = []
    for path in WORKSPACE_ROOT.rglob("*"):
        if len(results) >= max_results:
            break
        if not path.is_file():
            continue
        if mode == "name":
            if query.lower() in path.name.lower():
                results.append(str(path.relative_to(WORKSPACE_ROOT)))
        else:
            if path.suffix.lower() not in TEXT_EXTENSIONS:
                continue
            try:
                text = path.read_text(errors="ignore")
            except OSError:
                continue
            if query.lower() in text.lower():
                results.append(str(path.relative_to(WORKSPACE_ROOT)))
    return "\n".join(results) if results else "No matching files found."


@tool
def read_document(path: str) -> str:
    """Read and extract plain text from a file in the workspace.

    Supports .txt, .md, .pdf and .docx. Returns up to ~12000 characters;
    longer files are truncated with a notice.
    """
    try:
        full_path = _resolve(path)
    except ValueError as exc:
        return str(exc)
    if not full_path.exists():
        return f"File not found: {path}"

    suffix = full_path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        text = full_path.read_text(errors="ignore")
    elif suffix == ".pdf":
        reader = PdfReader(str(full_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        doc = DocxDocument(str(full_path))
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        return f"Unsupported file type '{suffix}'. Supported: .txt, .md, .pdf, .docx"

    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + f"\n\n[...truncated, {len(text) - MAX_CHARS} more characters...]"
    return text or "(file appears to be empty or text could not be extracted)"


@tool
def write_document(path: str, content: str) -> str:
    """Create or overwrite a document in the workspace with the given text content.

    The output format is chosen from the file extension: .docx produces a Word
    document (content split on blank lines into paragraphs), .txt/.md write plain text.
    """
    try:
        full_path = _resolve(path)
    except ValueError as exc:
        return str(exc)

    full_path.parent.mkdir(parents=True, exist_ok=True)
    suffix = full_path.suffix.lower()
    if suffix == ".docx":
        doc = DocxDocument()
        for paragraph in content.split("\n\n"):
            doc.add_paragraph(paragraph)
        doc.save(str(full_path))
    elif suffix in {".txt", ".md"}:
        full_path.write_text(content)
    else:
        return f"Unsupported output type '{suffix}'. Use .docx, .txt or .md"

    return f"Wrote {len(content)} characters to {full_path.relative_to(WORKSPACE_ROOT)}"


TOOLS = [search_files, read_document, write_document]